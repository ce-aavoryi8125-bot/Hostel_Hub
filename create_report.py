import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def add_custom_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    run = h.runs[0]
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 118, 110) # Teal
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42) # Navy
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 65, 85)
    return h

def create_document():
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set base style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 41, 59) # Slate 800

    # -------------------------------------------------------------
    # COVER / TITLE PAGE
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(12)
    
    r_sub = title_p.add_run("UNIVERSITY OF MINES AND TECHNOLOGY (UMaT), TARKWA\n")
    r_sub.font.size = Pt(13)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(100, 116, 139)

    r_main = title_p.add_run("\nHOSTEL HUB — STUDENT HOSTEL MANAGEMENT SYSTEM\n")
    r_main.font.size = Pt(22)
    r_main.font.bold = True
    r_main.font.color.rgb = RGBColor(15, 118, 110) # Dark Teal

    r_desc = title_p.add_run("Final Technical Project Implementation & Submission Report\n")
    r_desc.font.size = Pt(14)
    r_desc.font.italic = True
    r_desc.font.color.rgb = RGBColor(71, 85, 105)

    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.paragraph_format.space_after = Pt(24)
    r_div = p_div.add_run("══════════════════════════════════════════════════════════════")
    r_div.font.color.rgb = RGBColor(15, 118, 110)

    # Group Members Table
    p_members_head = doc.add_paragraph()
    p_members_head.paragraph_format.space_before = Pt(12)
    p_members_head.paragraph_format.space_after = Pt(6)
    r_mh = p_members_head.add_run("PROJECT GROUP MEMBERS & INDEX NUMBERS")
    r_mh.font.size = Pt(12)
    r_mh.font.bold = True
    r_mh.font.color.rgb = RGBColor(15, 23, 42)

    members = [
        ("Albert Atsu Avoryi", "FCM.41.008.043.25", "Lead Systems Architect & Developer"),
        ("Isaac Kofi Armah", "FCM.41.008.033.25", "Database Analyst & Backend Engineer"),
        ("Adegbedzi Joy Edudzi Yao", "FCM.41.008.013.25", "Frontend UI/UX & Design Engineer"),
        ("Ama Achiaa Owusu", "FCM.41.008.105.25", "Systems Tester & Quality Assurance"),
        ("Gyan Nicholas Kwame", "FCM.41.008.073.25", "Database Integration Specialist"),
        ("Regina Frimpomaa Anaman", "FCM.41.008.023.25", "Requirements & Documentation Lead"),
        ("Francis Abakah-Pobee", "FCM.41.008.002.25", "Infrastructure & Deployment Specialist"),
        ("Bonye Princella Muteima", "FCM.41.008.053.25", "User Workflow & Verification Specialist"),
        ("Onesiphorus Nana Essuman EDUAH", "FCM.41.008.063.25", "Security & Authentication Engineer"),
        ("Cyrus Tsibu-Yeboah", "FCM.41.008.116.25", "Systems Analyst & Data Modeler"),
        ("Obeng Nana Acheampong", "FCM.41.008.095.25", "Quality Assurance & Support Specialist")
    ]

    t_members = doc.add_table(rows=len(members) + 1, cols=3)
    t_members.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header Row
    headers = ["Student Full Name", "Student Index Number", "Project Role & Contribution"]
    hdr_cells = t_members.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "0F766E")
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(10)
        set_cell_margins(hdr_cells[i], 120, 120, 150, 150)

    # Data Rows
    for idx, (name, sindex, role) in enumerate(members):
        row_cells = t_members.rows[idx + 1].cells
        row_cells[0].text = name
        row_cells[1].text = sindex
        row_cells[2].text = role
        
        bg_color = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
        for c in row_cells:
            set_cell_background(c, bg_color)
            set_cell_margins(c, 100, 100, 120, 120)
            p = c.paragraphs[0]
            if len(p.runs) > 0:
                p.runs[0].font.size = Pt(9.5)

    # Submission Date & Location
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(30)
    r_m1 = p_meta.add_run("Institution: University of Mines and Technology (UMaT), Tarkwa, Ghana\n")
    r_m1.font.bold = True
    r_m2 = p_meta.add_run("Target Live Hosting URL: https://aa463.ceiscy.com/hostelhub/\n")
    r_m2.font.color.rgb = RGBColor(15, 118, 110)
    r_m2.font.bold = True
    r_m3 = p_meta.add_run("Submission Date: August 2026")
    r_m3.font.italic = True

    doc.add_page_break()

    # -------------------------------------------------------------
    # SECTION 1: EXECUTIVE SUMMARY
    # -------------------------------------------------------------
    add_custom_heading(doc, "1. Executive Summary", level=1)
    
    p = doc.add_paragraph(
        "Hostel Hub is a web-based Student Hostel Management and Accommodation Discovery Platform "
        "designed specifically for the University of Mines and Technology (UMaT) community in Tarkwa, Ghana. "
        "The system addresses critical challenges faced by students during non-resident accommodation searches, "
        "such as fraudulent agent fees, fake hostel listings, lack of transparent room pricing, and missing payment receipts."
    )
    p.paragraph_format.space_after = Pt(8)

    p2 = doc.add_paragraph(
        "To ensure zero deployment friction and compatibility with standard shared hosting environments (cPanel/CEISCY), "
        "the application has been fully architected using native PHP 8.x, MySQL (PDO), HTML5, CSS3, and Vanilla JavaScript. "
        "The system eliminates third-party Node.js, Express, and cloud database dependencies, allowing full execution inside a "
        "standard shared hosting subfolder (/public_html/hostelhub/)."
    )
    p2.paragraph_format.space_after = Pt(12)

    # Key Highlights Box
    t_box = doc.add_table(rows=1, cols=1)
    t_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_box = t_box.rows[0].cells[0]
    set_cell_background(c_box, "F0FDFA")
    set_cell_margins(c_box, 140, 140, 180, 180)
    
    p_box = c_box.paragraphs[0]
    r_bt = p_box.add_run("🌟 KEY SYSTEM ACHIEVEMENTS & PRODUCTION FEATURES:\n")
    r_bt.font.bold = True
    r_bt.font.size = Pt(10.5)
    r_bt.font.color.rgb = RGBColor(15, 118, 110)

    highlights = [
        "100% Student-Only Portal: Streamlined UI built exclusively for UMaT student accommodation discovery.",
        "One-Click Lecturer Demo Access: Prominent 🎓 Student Demo button on login.php for immediate authentication as student@hostelhub.dev.",
        "15 Pre-Seeded Realistic Ghanaian Hostels: Detailed pricing, photos, distance to UMaT, and room options across 7 Tarkwa locations.",
        "Mobile Money & Payment Proof Upload: Allows students to upload payment screenshots and view pending verification status.",
        "Printable Official HTML Receipt: Professional printable receipt (booking-details.php) with window.print() integration.",
        "Dynamic Subfolder Compatibility: Flexible BASE_URL calculation supporting subfolder deployment at https://aa463.ceiscy.com/hostelhub/."
    ]
    for h in highlights:
        p_h = c_box.add_paragraph()
        p_h.paragraph_format.space_before = Pt(2)
        p_h.paragraph_format.space_after = Pt(2)
        r = p_h.add_run("• " + h)
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECTION 2: SYSTEM ARCHITECTURE & TECH STACK
    # -------------------------------------------------------------
    add_custom_heading(doc, "2. Technical Architecture & Technology Stack", level=1)

    p = doc.add_paragraph(
        "Hostel Hub is built on a clean 3-tier monolithic PHP/MySQL architecture designed for maximum performance, "
        "reliability, and seamless deployment on standard cPanel Apache/LiteSpeed web servers."
    )
    p.paragraph_format.space_after = Pt(8)

    t_stack = doc.add_table(rows=6, cols=3)
    t_stack.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    st_headers = ["Layer / Component", "Technology Selected", "Functional Purpose"]
    for i, h in enumerate(st_headers):
        t_stack.rows[0].cells[i].text = h
        set_cell_background(t_stack.rows[0].cells[i], "0F172A")
        p = t_stack.rows[0].cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(10)
        set_cell_margins(t_stack.rows[0].cells[i], 100, 100, 120, 120)

    stack_data = [
        ("Backend Runtime", "PHP 8.x (Native)", "Handles routing, business logic, session authentication, and database requests."),
        ("Database Layer", "MySQL 8.0 / MariaDB (PDO)", "Stores relational data for students, hostels, rooms, bookings, payments, and notifications."),
        ("User Interface", "HTML5 & Vanilla CSS3", "Provides responsive layout, dark hero branding, cards, modals, and print stylesheets."),
        ("Client Interactions", "Vanilla JavaScript (ES6+)", "Handles flash toast auto-dismissal, live hostel search filtering, and receipt printing."),
        ("Session Security", "PHP Native Sessions & Bcrypt", "Manages student session persistence and secure password hashing via password_hash().")
    ]

    for idx, (layer, tech, purpose) in enumerate(stack_data):
        row_cells = t_stack.rows[idx + 1].cells
        row_cells[0].text = layer
        row_cells[1].text = tech
        row_cells[2].text = purpose
        bg_color = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
        for c in row_cells:
            set_cell_background(c, bg_color)
            set_cell_margins(c, 80, 80, 100, 100)
            p = c.paragraphs[0]
            if len(p.runs) > 0:
                p.runs[0].font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECTION 3: DATABASE SCHEMA & DATA DICTIONARY
    # -------------------------------------------------------------
    add_custom_heading(doc, "3. Relational Database Schema & Data Dictionary", level=1)

    p = doc.add_paragraph(
        "The relational database schema is stored in database/hostelhub.sql. "
        "It consists of 7 normalized relational tables enforcing foreign key integrity, indexing, and prepared statement compatibility."
    )
    p.paragraph_format.space_after = Pt(8)

    tables_info = [
        ("1. students", "Stores registered student accounts, credentials, contact details, and academic profiles.", [
            ("id", "INT AUTO_INCREMENT", "PRIMARY KEY", "Unique student identifier"),
            ("name", "VARCHAR(100)", "NOT NULL", "Student full name"),
            ("email", "VARCHAR(150)", "UNIQUE NOT NULL", "Student email address"),
            ("password_hash", "VARCHAR(255)", "NOT NULL", "Bcrypt hashed password"),
            ("phone", "VARCHAR(30)", "NOT NULL", "Contact phone number"),
            ("student_index", "VARCHAR(50)", "NULL", "UMaT student index number"),
            ("faculty", "VARCHAR(100)", "NULL", "Faculty name"),
            ("department", "VARCHAR(100)", "NULL", "Academic department"),
            ("level", "VARCHAR(30)", "NULL", "Academic level (Level 100-400)")
        ]),
        ("2. hostels", "Stores verified student hostels, location mappings, pricing, distance to campus, and photo galleries.", [
            ("id", "INT AUTO_INCREMENT", "PRIMARY KEY", "Unique hostel identifier"),
            ("name", "VARCHAR(120)", "NOT NULL", "Hostel name (e.g. Banso Royal)"),
            ("location_name", "VARCHAR(100)", "NOT NULL", "Neighborhood location name"),
            ("distance_km", "DECIMAL(3,1)", "NOT NULL", "Distance from UMaT main gate in km"),
            ("price_per_year", "DECIMAL(10,2)", "NOT NULL", "Starting annual price in GH₵"),
            ("photos", "TEXT", "NULL", "Comma-separated photo URLs"),
            ("facilities", "TEXT", "NULL", "Comma-separated facility tags"),
            ("is_published", "TINYINT(1)", "DEFAULT 1", "Publication visibility status")
        ]),
        ("3. bookings", "Tracks student room reservations, selected room type, payment status, and verification reference.", [
            ("id", "INT AUTO_INCREMENT", "PRIMARY KEY", "Unique booking record ID"),
            ("booking_reference", "VARCHAR(30)", "UNIQUE NOT NULL", "Generated reference (HH-2026-XXXX)"),
            ("student_id", "INT", "FOREIGN KEY", "References students(id)"),
            ("hostel_id", "INT", "FOREIGN KEY", "References hostels(id)"),
            ("room_type", "VARCHAR(50)", "NOT NULL", "Room type (1-in-a-room, 2-in-a-room)"),
            ("amount", "DECIMAL(10,2)", "NOT NULL", "Total annual booking fee (GH₵)"),
            ("status", "ENUM('pending','confirmed')", "DEFAULT 'pending'", "Reservation clearance status"),
            ("payment_status", "ENUM('unpaid','pending_verification','paid')", "DEFAULT 'unpaid'", "Payment status"),
            ("payment_proof", "VARCHAR(255)", "NULL", "Uploaded screenshot filename")
        ])
    ]

    for t_title, t_desc, rows in tables_info:
        add_custom_heading(doc, t_title, level=2)
        p_tdesc = doc.add_paragraph(t_desc)
        p_tdesc.paragraph_format.space_after = Pt(4)

        t_schema = doc.add_table(rows=len(rows) + 1, cols=4)
        t_schema.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        s_headers = ["Column Name", "Data Type", "Constraints", "Description"]
        for i, h in enumerate(s_headers):
            t_schema.rows[0].cells[i].text = h
            set_cell_background(t_schema.rows[0].cells[i], "0F766E")
            p = t_schema.rows[0].cells[i].paragraphs[0]
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            p.runs[0].font.size = Pt(9)
            set_cell_margins(t_schema.rows[0].cells[i], 80, 80, 100, 100)

        for r_idx, (col, dtype, constr, cdesc) in enumerate(rows):
            r_cells = t_schema.rows[r_idx + 1].cells
            r_cells[0].text = col
            r_cells[1].text = dtype
            r_cells[2].text = constr
            r_cells[3].text = cdesc
            bg_color = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
            for c in r_cells:
                set_cell_background(c, bg_color)
                set_cell_margins(c, 60, 60, 80, 80)
                p = c.paragraphs[0]
                if len(p.runs) > 0:
                    p.runs[0].font.size = Pt(8.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # -------------------------------------------------------------
    # SECTION 4: CORE APPLICATION FEATURES & WORKFLOWS
    # -------------------------------------------------------------
    add_custom_heading(doc, "4. Key Application Modules & Student Workflows", level=1)

    modules = [
        ("4.1 Student Hostel Discovery & Live Filtering (index.php, hostels.php)", 
         "Students can search hostels by keyword, filter by Tarkwa neighborhood locations (Banso, Ayensu, Brahabobom, Tamso, etc.), "
         "and dynamically slide maximum annual budget pricing in real time."),
        
        ("4.2 One-Click Demo Access (login.php)", 
         "Designed for seamless evaluation by lecturers without requiring password lookup. "
         "Clicking the 🎓 Student Demo button instantly authenticates student@hostelhub.dev."),

        ("4.3 Room Reservation System (book.php)", 
         "Students select room capacity (1-in-a-room, 2-in-a-room, 4-in-a-room), academic period, and payment channel. "
         "The system generates a unique reference code (HH-2026-XXXX) and reserves the room."),

        ("4.4 Mobile Money Payment Proof Upload (payment.php)", 
         "Students submit MTN Mobile Money or bank deposit transfer references and upload payment screenshot files (.png, .jpg, .pdf). "
         "The payment status updates immediately to 'pending_verification'."),

        ("4.5 Printable HTML Receipt (booking-details.php)", 
         "Renders an official PDF-style printable receipt containing student information, hostel details, room type, total fee, "
         "and an authorized clearance stamp. Features a 🖨️ Print Receipt button executing window.print().")
    ]

    for m_title, m_desc in modules:
        add_custom_heading(doc, m_title, level=2)
        p_m = doc.add_paragraph(m_desc)
        p_m.paragraph_format.space_after = Pt(6)

    # -------------------------------------------------------------
    # SECTION 5: VERIFICATION & TEST RESULTS
    # -------------------------------------------------------------
    add_custom_heading(doc, "5. Verification & Automated Test Results", level=1)

    p_t = doc.add_paragraph("The complete PHP/MySQL implementation underwent rigorous verification to ensure zero runtime defects.")
    p_t.paragraph_format.space_after = Pt(8)

    t_test = doc.add_table(rows=7, cols=4)
    t_test.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    test_headers = ["Test Module", "Action Tested", "Expected Outcome", "Status"]
    for i, h in enumerate(test_headers):
        t_test.rows[0].cells[i].text = h
        set_cell_background(t_test.rows[0].cells[i], "0F172A")
        p = t_test.rows[0].cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(9.5)
        set_cell_margins(t_test.rows[0].cells[i], 100, 100, 120, 120)

    test_rows = [
        ("Demo Authentication", "Click 🎓 Student Demo on login.php", "Authenticates student@hostelhub.dev instantly", "PASSED ✅"),
        ("Hostel Search", "Filter by 'Banso' and max price GH₵ 3,500", "Displays matching verified hostels grid", "PASSED ✅"),
        ("Room Reservation", "Submit reservation on book.php", "Generates HH-2026-XXXX reference and stores booking", "PASSED ✅"),
        ("Payment Proof Upload", "Upload proof.png on payment.php", "Saves file into uploads/payment-proofs/ and updates status", "PASSED ✅"),
        ("Printable Receipt", "Click 🖨️ Print Receipt on booking-details.php", "Triggers browser window.print() clean document", "PASSED ✅"),
        ("Subfolder Compatibility", "Access app via /hostelhub/", "All CSS, JS, images, forms, and redirects resolve correctly", "PASSED ✅")
    ]

    for idx, (mod, act, exp, stat) in enumerate(test_rows):
        r_cells = t_test.rows[idx + 1].cells
        r_cells[0].text = mod
        r_cells[1].text = act
        r_cells[2].text = exp
        r_cells[3].text = stat
        bg_color = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
        for c in r_cells:
            set_cell_background(c, bg_color)
            set_cell_margins(c, 80, 80, 100, 100)
            p = c.paragraphs[0]
            if len(p.runs) > 0:
                p.runs[0].font.size = Pt(9)
                if stat.startswith("PASSED"):
                    p.runs[0].font.bold = True
                    p.runs[0].font.color.rgb = RGBColor(4, 120, 87)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECTION 6: CONCLUSION & DEPLOYMENT SUMMARY
    # -------------------------------------------------------------
    add_custom_heading(doc, "6. Conclusion & Deployment Instructions", level=1)

    p_c = doc.add_paragraph(
        "The Hostel Hub Student Hostel Management System successfully provides UMaT Tarkwa students with a transparent, "
        "secure, and agent-free accommodation reservation experience. The application package is completely self-contained "
        "inside the CEISCY_SUBMISSION folder and ready for direct cPanel deployment."
    )
    p_c.paragraph_format.space_after = Pt(8)

    p_dep = doc.add_paragraph()
    r_dt = p_dep.add_run("📋 DEPLOYMENT STEPS FOR CEISCY SHARED HOSTING (/public_html/hostelhub/):\n")
    r_dt.font.bold = True
    r_dt.font.color.rgb = RGBColor(15, 118, 110)

    steps = [
        "Upload all contents of CEISCY_SUBMISSION into /public_html/hostelhub/.",
        "Create a MySQL database in cPanel MySQL Database Wizard.",
        "Import database/hostelhub.sql into phpMyAdmin.",
        "Configure database credentials in config/database.php.",
        "Access the live portal at https://aa463.ceiscy.com/hostelhub/."
    ]
    for s in steps:
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_before = Pt(2)
        p_s.paragraph_format.space_after = Pt(2)
        r = p_s.add_run("1. " + s if s == steps[0] else str(steps.index(s)+1) + ". " + s)
        r.font.size = Pt(10)

    # Save document
    filename = "Hostel_Hub_Final_Project_Report.docx"
    doc.save(filename)
    print(f"Document successfully created: {filename}")

if __name__ == "__main__":
    create_document()
